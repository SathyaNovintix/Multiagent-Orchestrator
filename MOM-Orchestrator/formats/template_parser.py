"""
Template Parser — AgentMesh AI
Dynamically parses uploaded Excel/Word template files to extract ANY structure.
No hardcoded section names - fully adaptive to user's template.
"""
from __future__ import annotations
import io
import re
from pathlib import Path
from typing import List, Dict, Optional


def is_likely_header(text: str, cell_style=None) -> bool:
    """
    Heuristic to detect if a cell/paragraph is likely a section header.
    Checks for: uppercase, short length, ends with colon, bold formatting.
    """
    if not text or len(text.strip()) < 2:
        return False
    
    text = text.strip()
    
    # Exclude common field names that shouldn't be section headers
    excluded_patterns = [
        r'^s\.?\s*no\.?$',  # S.No, S No, SNo
        r'^sr\.?\s*no\.?$',  # Sr.No
        r'^#$',
        r'^\d+$',  # Pure numbers
        r'^\d{1,2}/\d{1,2}/\d{2,4}$',  # Dates like 3/23/2026
        r'^[A-Za-z]{3}\s\d{1,2},\s\d{4}$',  # Dates like March 30, 2026
    ]
    for pattern in excluded_patterns:
        if re.match(pattern, text, re.IGNORECASE):
            return False
    
    # Strong indicators
    if text.isupper() and len(text) < 50 and len(text) > 3:
        return True
    if text.endswith(':') and len(text) < 50:
        return True
    if cell_style and hasattr(cell_style, 'font') and cell_style.font.bold:
        return True
    
    # Moderate indicators: starts with number/letter followed by period or parenthesis
    if re.match(r'^[A-Z0-9]+[\.\)]\s+[A-Z]', text):
        return True
    
    # Title case and short (but not too short)
    if text.istitle() and 3 <= len(text.split()) <= 5 and len(text) < 50:
        return True
    
    return False


def extract_fields_from_row(row: tuple) -> List[Dict[str, any]]:
    """
    Extract field definitions from a row (typically the row after a section header).
    Returns list of field definitions with name and column index.
    """
    fields = []
    # Identify placeholder patterns that shouldn't be fields
    field_blacklist = [r'^row\s*\d+$', r'^column\s*\d+$', r'^field\s*\d+$']
    
    for idx, cell in enumerate(row):
        if cell and isinstance(cell, str) and len(cell.strip()) > 0:
            cell_text = cell.strip()
            # Skip very long text (likely data, not a field name)
            if len(cell_text) < 100 and not cell_text.startswith('='):
                # Skip if it matches a blacklist pattern
                if any(re.match(p, cell_text, re.IGNORECASE) for p in field_blacklist):
                    continue
                fields.append({
                    'name': cell_text,
                    'column': idx
                })
    return fields


def parse_excel_template(file_bytes: bytes, filename: str) -> dict:
    """
    Dynamically parses Excel/CSV file structure.
    Detects section headers and field columns without hardcoded keywords.
    """
    import openpyxl
    
    # Check if it's a CSV file
    if filename.lower().endswith('.csv'):
        import csv
        import io as io_module
        
        # Try to parse as CSV
        text_content = file_bytes.decode('utf-8', errors='ignore')
        csv_reader = csv.reader(io_module.StringIO(text_content))
        rows_list = list(csv_reader)
        
        sections = []
        all_fields = {}
        
        # Strategy: Detect structure intelligently
        # 1. First row with single cell = document title
        # 2. Rows with 2 cells (label: value) = metadata fields
        # 3. Rows with uppercase/bold text = section headers
        # 4. Rows with multiple columns after a header = table sections
        
        i = 0
        while i < len(rows_list):
            row = rows_list[i]
            
            if not row or not any(row):
                i += 1
                continue
            
            # Get non-empty cells
            non_empty = [cell for cell in row if cell and str(cell).strip()]
            
            if not non_empty:
                i += 1
                continue
            
            first_cell = str(row[0]).strip() if row[0] else ""
            
            # Skip pure number rows (likely data rows) - but keep scanning for more headers
            if first_cell.isdigit() and len(first_cell) <= 3:
                i += 1
                continue
            
            # Skip S.No rows only if they're standalone (not part of a table header)
            if re.match(r'^s\.?\s*no\.?$', first_cell, re.IGNORECASE) and len(non_empty) == 1:
                i += 1
                continue
            
            # Case 1: Single cell with uppercase text = document title or section header
            # This should work for both "MINUTES OF MEETING" and "Action Points"
            if len(non_empty) == 1 and len(first_cell) > 3:
                # Check if it's a header (uppercase, title case, or ends with colon)
                is_header_candidate = (
                    first_cell.isupper() or 
                    is_likely_header(first_cell) or
                    (first_cell.istitle() and len(first_cell.split()) <= 4)
                )
                
                if is_header_candidate:
                    section_id = re.sub(r'[^a-z0-9]+', '_', first_cell.lower()).strip('_')
                    if not section_id:
                        section_id = f"section_{len(sections)}"
                    
                    # Ensure unique ID
                    existing_ids = [s['id'] for s in sections]
                    if section_id in existing_ids:
                        section_id = f"{section_id}_{len(sections)}"
                    
                    sections.append({
                        'id': section_id,
                        'label': first_cell,
                        'row': i,
                        'type': 'header'
                    })
                    
                    # Check if next row has column headers (for tables)
                    # Skip up to 15 empty rows to find the table header
                    check_row_idx = i + 1
                    found_header = False
                    while check_row_idx < len(rows_list) and check_row_idx < i + 20:
                        next_row = rows_list[check_row_idx]
                        next_non_empty = [cell for cell in next_row if cell and str(cell).strip()]
                        
                        if len(next_non_empty) > 1:
                            # Found table headers
                            all_fields[section_id] = [
                                {
                                    'id': re.sub(r'[^a-z0-9]+', '_', str(cell).strip().lower()).strip('_') or f'col_{idx}',
                                    'name': str(cell).strip(),
                                    'label': str(cell).strip(),
                                    'type': 'text',
                                    'column': idx
                                }
                                for idx, cell in enumerate(next_row)
                                if cell and str(cell).strip()
                            ]
                            found_header = True
                            break
                        
                        check_row_idx += 1
                    
                    if found_header:
                        i = check_row_idx  # Skip to the header row
            
            # Case 2: Two cells (key: value) = metadata field
            elif len(non_empty) == 2 and len(row) >= 2:
                label = first_cell
                value = str(row[1]).strip() if len(row) > 1 and row[1] else ""
                
                # Only treat as metadata if the label looks like a field name (not a number)
                if not first_cell.isdigit() and len(first_cell) > 1:
                    field_id = re.sub(r'[^a-z0-9]+', '_', label.lower()).strip('_')
                    if not field_id:
                        field_id = f"field_{len(all_fields)}"
                    
                    # If previous section was a header, add this as a field to it
                    if sections and sections[-1].get('type') == 'header':
                        prev_section_id = sections[-1]['id']
                        if prev_section_id not in all_fields:
                            all_fields[prev_section_id] = []
                        
                        # Ensure field ID is unique within the section
                        existing_fids = [f['id'] for f in all_fields[prev_section_id]]
                        if field_id in existing_fids:
                            field_id = f"{field_id}_{len(existing_fids)}"
                            
                        all_fields[prev_section_id].append({
                            'id': field_id,
                            'name': label,
                            'label': label,
                            'type': 'text',
                            'column': 0
                        })
                    else:
                        # No parent header, create a standalone metadata section
                        section_id = field_id
                        existing_ids = [s['id'] for s in sections]
                        if section_id in existing_ids:
                            section_id = f"{section_id}_{len(sections)}"
                        
                        sections.append({
                            'id': section_id,
                            'label': label,
                            'row': i,
                            'type': 'metadata'
                        })
                        all_fields[section_id] = [{
                            'id': section_id,
                            'name': label,
                            'label': label,
                            'type': 'text',
                            'column': 0
                        }]
            
            # Case 3: Multiple cells = potential table header or data row
            elif len(non_empty) > 2:
                # Check if this looks like a header row (not all numbers)
                non_numeric = [cell for cell in non_empty if not str(cell).strip().isdigit()]
                if len(non_numeric) >= len(non_empty) * 0.5:  # At least 50% non-numeric
                    # Could be a section header or table header
                    # If previous section exists and has no fields, this might be its column headers
                    if sections and sections[-1].get('type') == 'header' and sections[-1]['id'] not in all_fields:
                        # Add as fields to the previous section
                        prev_section_id = sections[-1]['id']
                        all_fields[prev_section_id] = [
                            {
                                'id': re.sub(r'[^a-z0-9]+', '_', str(cell).strip().lower()).strip('_') or f'col_{idx}',
                                'name': str(cell).strip(),
                                'label': str(cell).strip(),
                                'type': 'text',
                                'column': idx
                            }
                            for idx, cell in enumerate(row)
                            if cell and str(cell).strip()
                        ]
            
            i += 1
        
        return {
            'sections': sections,
            'fields': all_fields,
            'source': 'csv',
            'filename': filename,
            'total_sections': len(sections),
            'note': 'Intelligently parsed CSV structure'
        }
    
    # Otherwise parse as Excel
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    sheet = wb.active
    
    sections = []
    all_fields = {}
    
    rows_list = []
    for row in sheet.iter_rows(min_row=1, max_row=200, values_only=False):
        row_values = [cell.value for cell in row]
        row_objs = [cell for cell in row]
        rows_list.append((row_values, row_objs))
    
    i = 0
    while i < len(rows_list):
        vals, objs = rows_list[i]
        
        if not vals or not any(v is not None for v in vals):
            i += 1
            continue
            
        non_empty = [v for v in vals if v is not None and str(v).strip()]
        if not non_empty:
            i += 1
            continue
            
        first_val = vals[0]
        first_text = str(first_val).strip() if first_val is not None else ""
        first_obj = objs[0]
        
        # Case 1: Likely a Section Header
        if is_likely_header(first_text, first_obj) and len(non_empty) == 1:
            section_id = re.sub(r'[^a-z0-9]+', '_', first_text.lower()).strip('_')
            if not section_id:
                section_id = f"section_{len(sections)}"
            
            # Ensure unique ID
            existing_ids = [s['id'] for s in sections]
            if section_id in existing_ids:
                section_id = f"{section_id}_{len(sections)}"
            
            sections.append({
                'id': section_id,
                'label': first_text,
                'row': i,
                'type': 'header'
            })
            
            # Scan for fields/columns for this section
            check_idx = i + 1
            found_fields = False
            while check_idx < len(rows_list) and check_idx < i + 15:
                n_vals, n_objs = rows_list[check_idx]
                n_non_empty = [v for v in n_vals if v is not None and str(v).strip()]
                
                if len(n_non_empty) > 1:
                    # Found column headers
                    all_fields[section_id] = [
                        {
                            'id': re.sub(r'[^a-z0-9]+', '_', str(v).strip().lower()).strip('_') or f'col_{idx}',
                            'name': str(v).strip(),
                            'label': str(v).strip(),
                            'type': 'text',
                            'column': idx
                        }
                        for idx, v in enumerate(n_vals)
                        if v is not None and str(v).strip()
                    ]
                    found_fields = True
                    break
                check_idx += 1
            
            if found_fields:
                i = check_idx
        
        # Case 2: Key-Value Metadata field
        elif len(non_empty) == 2:
            label = str(non_empty[0]).strip()
            field_id = re.sub(r'[^a-z0-9]+', '_', label.lower()).strip('_')
            
            # Exclude dates/numbers from being field labels
            is_valid_label = (
                not label.isdigit() and 
                len(label) > 1 and 
                not re.match(r'^\d{1,2}/\d{1,2}/\d{2,4}$', label) and
                not re.match(r'^[A-Za-z]{3}\s\d{1,2},\s\d{4}$', label)
            )
            
            if is_valid_label:
                # Group under previous header if it exists
                if sections and sections[-1].get('type') == 'header':
                    prev_id = sections[-1]['id']
                    if prev_id not in all_fields:
                        all_fields[prev_id] = []
                    
                    if not any(f['name'] == label for f in all_fields[prev_id]):
                        all_fields[prev_id].append({
                            'id': field_id,
                            'name': label,
                            'label': label,
                            'type': 'text',
                            'column': 0
                        })
                else:
                    # Standalone metadata row
                    sect_id = field_id
                    if any(s['id'] == sect_id for s in sections):
                        sect_id = f"{sect_id}_{len(sections)}"
                    
                    sections.append({
                        'id': sect_id,
                        'label': label,
                        'row': i,
                        'type': 'metadata'
                    })
                    all_fields[sect_id] = [{
                        'id': sect_id,
                        'name': label,
                        'label': label,
                        'type': 'text',
                        'column': 0
                    }]
        
        # Case 3: Multiple cells = potential table header or data row
        elif len(non_empty) > 2:
            # Check if this looks like a header row (not all numbers)
            non_numeric = [v for v in non_empty if not str(v).strip().isdigit()]
            if len(non_numeric) >= len(non_empty) * 0.5:
                # If previous section exists and has no fields, this might be its column headers
                if sections and sections[-1].get('type') == 'header' and sections[-1]['id'] not in all_fields:
                    prev_id = sections[-1]['id']
                    all_fields[prev_id] = [
                        {
                            'id': re.sub(r'[^a-z0-9]+', '_', str(v).strip().lower()).strip('_') or f'col_{idx}',
                            'name': str(v).strip(),
                            'label': str(v).strip(),
                            'type': 'text',
                            'column': idx
                        }
                        for idx, v in enumerate(vals)
                        if v is not None and str(v).strip()
                    ]
        i += 1


    return {
        'sections': sections,
        'fields': all_fields,
        'source': 'excel',
        'filename': filename,
        'total_sections': len(sections)
    }



def parse_word_template(file_bytes: bytes, filename: str) -> dict:
    """
    Dynamically parses Word document structure.
    Detects any heading-style text as sections without hardcoded keywords.
    """
    from docx import Document
    
    doc = Document(io.BytesIO(file_bytes))
    
    sections = []
    all_fields = {}
    
    # Extract sections from paragraphs with heading styles or bold text
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        if not text or len(text) < 2:
            continue
        
        # Check if it's a heading (style-based or bold)
        is_heading_style = para.style.name.startswith('Heading')
        is_bold = para.runs and len(para.runs) > 0 and para.runs[0].bold
        
        # Also check heuristics
        is_header = is_heading_style or is_bold or is_likely_header(text)
        
        if is_header:
            # Create section ID from text
            section_id = re.sub(r'[^a-z0-9]+', '_', text.lower()).strip('_')
            if not section_id:
                section_id = f"section_{len(sections) + 1}"
            
            # Ensure unique ID
            existing_ids = [s['id'] for s in sections]
            if section_id in existing_ids:
                section_id = f"{section_id}_{len(sections) + 1}"
            
            sections.append({
                'id': section_id,
                'label': text,
                'paragraph_index': para_idx
            })
    
    # Extract sections from tables (each table becomes a section)
    for table_idx, table in enumerate(doc.tables):
        if table.rows and len(table.rows) > 0:
            # Use first row as potential section name or generate one
            first_row_text = ' '.join(cell.text.strip() for cell in table.rows[0].cells if cell.text.strip())
            
            if first_row_text and len(first_row_text) < 100:
                section_label = first_row_text
            else:
                section_label = f"Table {table_idx + 1}"
            
            section_id = re.sub(r'[^a-z0-9]+', '_', section_label.lower()).strip('_')
            if not section_id:
                section_id = f"table_{table_idx + 1}"
            
            # Ensure unique ID
            existing_ids = [s['id'] for s in sections]
            if section_id in existing_ids:
                section_id = f"{section_id}_{table_idx + 1}"
            
            sections.append({
                'id': section_id,
                'label': section_label,
                'table_index': table_idx
            })
            
            # Extract field names from table header row
            if len(table.rows) > 1:
                header_cells = [cell.text.strip() for cell in table.rows[0].cells]
                fields = []
                for col_idx, header in enumerate(header_cells):
                    if header and len(header) < 100:
                        fields.append({
                            'name': header,
                            'column': col_idx
                        })
                if fields:
                    all_fields[section_id] = fields
    
    return {
        'sections': sections,
        'fields': all_fields,
        'source': 'word',
        'filename': filename,
        'total_sections': len(sections)
    }


def parse_template_file(file_bytes: bytes, filename: str) -> dict:
    """
    Main entry point for parsing template files.
    Auto-detects format and dynamically extracts structure.
    """
    ext = Path(filename).suffix.lower()
    
    try:
        if ext in ['.xlsx', '.xls', '.csv']:
            result = parse_excel_template(file_bytes, filename)
        elif ext in ['.docx', '.doc']:
            result = parse_word_template(file_bytes, filename)
        else:
            raise ValueError(f"Unsupported file format: {ext}. Supported: .xlsx, .xls, .csv, .docx, .doc")
        
        # If no sections were detected, return empty structure (let AI decide)
        if not result.get('sections'):
            result['sections'] = []
            result['note'] = "No clear structure detected. AI will generate standard MOM format."
        
        return result
        
    except ImportError as e:
        # Missing library
        return {
            'sections': [],
            'fields': {},
            'source': 'error',
            'filename': filename,
            'parse_error': f"Missing required library: {str(e)}. Install openpyxl or python-docx."
        }
    except Exception as e:
        # Return error info but don't fail
        return {
            'sections': [],
            'fields': {},
            'source': 'error',
            'filename': filename,
            'parse_error': str(e)
        }
